import os
import re
import uuid
import base64
import zipfile
import requests
from urllib.parse import urlparse
from xml.etree import ElementTree as ET
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

# =========================
# 1. 路径配置
# =========================
desktop      = os.path.join(os.path.expanduser("~"), "Desktop")
titles_path  = os.path.join(desktop, "title_BIT.txt")
download_dir = os.path.join(desktop, "BIT_原始文件")
output_path  = os.path.join(desktop, "title_BIT_output.xlsx")

os.makedirs(download_dir, exist_ok=True)

# =========================
# 2. AES-GCM 解密（密钥仅通过环境变量 BITMAIN_TITLE_ENCRYPT_AES_KEY_B64 提供）
# =========================
def _aes_key_from_env(var_name: str) -> bytes:
    raw = os.environ.get(var_name, "").strip()
    if not raw:
        return b""
    try:
        return base64.b64decode(raw)
    except Exception:
        print(f"解密失败: 环境变量 {var_name} 不是有效的 Base64")
        return b""


AES_KEY = _aes_key_from_env("BITMAIN_TITLE_ENCRYPT_AES_KEY_B64")


def decrypt_title(hex_str: str) -> str:
    """从 HEX 字符串解密 AES-GCM 密文，返回明文标题。"""
    if len(AES_KEY) not in (16, 24, 32):
        return ""
    try:
        raw = bytes.fromhex(hex_str)
        nonce = raw[:12]
        ciphertext_tag = raw[12:]
        plaintext = AESGCM(AES_KEY).decrypt(nonce, ciphertext_tag, None)
        return plaintext.decode("utf-8")
    except Exception as e:
        print(f"解密失败: {e}")
        return ""


def sanitize_filename(name: str) -> str:
    """移除文件名中不合法的字符。"""
    return re.sub(r'[\\/:*?"<>|]', '_', name).strip()


def normalize_plain_title(title: str) -> str:
    """规范化明文标题：仅保留文件名主体（去前缀和后缀）。"""
    if not title:
        return ""
    normalized = title.strip()
    normalized = normalized.replace("\\", "/").split("/")[-1].strip()
    colon_index = max(normalized.rfind(":"), normalized.rfind("："))
    if colon_index >= 0 and colon_index + 1 < len(normalized):
        normalized = normalized[colon_index + 1:].strip()
    removable_exts = {
        ".pdf", ".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx",
        ".txt", ".md", ".rtf", ".csv", ".zip"
    }
    root, ext = os.path.splitext(normalized)
    if ext and ext.lower() in removable_exts:
        normalized = root.strip()
    return normalized


# =========================
# 3. 读取 title_PR.txt
#    格式：前4行为列名（id / url / title / owner），
#    之后每行为 tab 分隔的数据行：id\turl\ttitle_hex\towner
# =========================
rows = []   # list of (url, title_hex, owner)

with open(titles_path, encoding="utf-8-sig") as f:
    lines = f.readlines()

for line in lines[4:]:
    parts = line.rstrip("\n").split("\t")
    if len(parts) < 4:
        continue
    url       = parts[1].strip()
    title_hex = parts[2].strip()
    owner     = parts[3].strip()
    if not url.startswith("http"):
        continue
    rows.append((url, title_hex, owner))

print(f"共读取 {len(rows)} 条记录")

# =========================
# 4. Content-Type 映射 & magic number 识别
# =========================
CONTENT_TYPE_MAP = {
    "application/pdf": ".pdf",
    "application/msword": ".doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
}

def detect_ext_by_magic(content: bytes) -> str:
    if content.startswith(b"%PDF"):
        return ".pdf"
    if content.startswith(b"\xD0\xCF\x11\xE0"):
        return ".doc"
    if content.startswith(b"PK"):
        return ".docx"
    return ".bin"

def resolve_unique_path(save_dir: str, file_name: str):
    base, ext = os.path.splitext(file_name)
    candidate = os.path.join(save_dir, file_name)
    if not os.path.exists(candidate):
        return candidate, False
    index = 2
    while True:
        new_name = f"{base} ({index}){ext}"
        candidate = os.path.join(save_dir, new_name)
        if not os.path.exists(candidate):
            return candidate, True
        index += 1

# =========================
# 5. 下载函数
# =========================
def download_file(url: str, save_dir: str, title: str = "") -> str | None:
    """下载文件到 save_dir，返回保存的完整路径；失败返回 None。"""
    try:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        content = resp.content
        if len(content) == 0:
            print(f"⚠️ 跳过空文件: {url}")
            return None
        parsed = urlparse(url)
        url_filename = os.path.basename(parsed.path)
        _, ext = os.path.splitext(url_filename)
        if not ext:
            content_type = resp.headers.get("Content-Type", "").split(";")[0]
            ext = CONTENT_TYPE_MAP.get(content_type, "")
        if not ext:
            ext = detect_ext_by_magic(content)
        name = sanitize_filename(title) if title else ""
        if not name:
            name = str(uuid.uuid4())
        final_name = f"{name}{ext}"
        file_path, renamed = resolve_unique_path(save_dir, final_name)
        with open(file_path, "wb") as f:
            f.write(content)
        label = "（重名已改名）" if renamed else ""
        print(f"✅ 下载成功{label}: {os.path.basename(file_path)} ({len(content) // 1024} KB)")
        return file_path
    except Exception as e:
        print(f"❌ 下载失败: {url}\n   原因: {e}")
        return None


# =========================
# 6. 在已下载文件夹中定位文件
# =========================
def find_downloaded_file(title: str, download_dir: str) -> str | None:
    """
    根据解密标题在 download_dir 中查找对应文件。
    匹配逻辑：文件名去后缀后以 sanitize 后的标题开头（兼容重名改名的情况）。
    """
    if not title:
        return None
    safe_title = sanitize_filename(title)
    if not safe_title:
        return None
    try:
        for fname in os.listdir(download_dir):
            name_no_ext, _ = os.path.splitext(fname)
            # 去掉末尾可能的 " (2)" " (3)" 冲突后缀再比对
            name_base = re.sub(r'\s+\(\d+\)$', '', name_no_ext).strip()
            if name_base == safe_title:
                return os.path.join(download_dir, fname)
    except Exception as e:
        print(f"⚠️ 查找文件失败: {e}")
    return None


# =========================
# 7. 提取管线：修订记录表 → OCR
#    无法提取时根据文件内容状态给出明确备注
#    状态: "empty"(内容为空) | "no_revision"(无修订记录) | "ok"(提取成功) | "null"(提取失败)
# =========================

_VERSION_KEYWORDS = ['版本', 'Version', 'version']
_AUTHOR_KEYWORDS  = ['作者', '编写', '拟制', '修订人', '编制', 'Author', 'author']

# 单页封面判定：去除空白后字符数低于此阈值视为"仅封面页"
_COVER_PAGE_CHAR_THRESHOLD = 800


def _get_unique_row_cells(row) -> list[str]:
    """获取行中去除合并单元格重复值后的文本列表（处理 Word 合并单元格）。"""
    seen_ids = set()
    texts = []
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id not in seen_ids:
            seen_ids.add(cell_id)
            texts.append(cell.text.strip())
    return texts


def _parse_table_data(table_data: list[list[str | None]]) -> tuple[str, str]:
    """
    从二维表格数据（第 0 行为表头）中识别版本列和作者列，
    返回最后一个非空数据行的 (author, version)。
    """
    if not table_data or len(table_data) < 2:
        return "", ""

    header = [str(c or "").strip() for c in table_data[0]]
    version_col = author_col = -1
    for j, h in enumerate(header):
        if version_col < 0 and any(k in h for k in _VERSION_KEYWORDS):
            version_col = j
        if author_col < 0 and any(k in h for k in _AUTHOR_KEYWORDS):
            author_col = j

    last_data = None
    for row in reversed(table_data[1:]):
        cells = [str(c or "").strip() for c in row]
        if any(cells):
            last_data = cells
            break

    if last_data is None:
        return "", ""

    version = last_data[version_col] if 0 <= version_col < len(last_data) else ""
    author  = last_data[author_col]  if 0 <= author_col  < len(last_data) else ""
    # 单元格内换行（多人合著）转为顿号，并去除多余空白
    author  = re.sub(r'\s*\n\s*', '、', author).strip()
    version = re.sub(r'\s*\n\s*', ' ', version).strip()
    return author, version


# ------------------------------------------------------------------
# 7a. docx 综合分析
# ------------------------------------------------------------------
def _analyze_docx(file_path: str) -> tuple[str, str, str]:
    """
    返回 (status, author, version)
      "empty"       — 文件无实质内容（空文件或仅封面页）
      "no_revision" — 有内容但未找到修订记录章节
      "ok"          — 成功提取
      "null"        — 找到修订记录但无法提取作者/版本
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table as DocxTable
    except ImportError:
        return "null", "", ""

    try:
        doc = Document(file_path)

        # ── 步骤1：计算文档总文字量 ──────────────────────────────
        all_chars = re.sub(r'\s', '', ''.join(
            p.text for p in doc.paragraphs
        ) + ''.join(
            cell.text for tbl in doc.tables
            for row in tbl.rows for cell in row.cells
        ))

        if len(all_chars) == 0:
            return "empty", "", ""

        # ── 步骤2：查找"修订记录"标题 + 后续表格（优先提取，成功则免判封面页）──
        found_heading = False
        for elem in doc.element.body:
            if elem.tag == qn('w:p'):
                text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
                if '修订记录' in text:
                    found_heading = True
            elif elem.tag == qn('w:tbl') and found_heading:
                tbl = DocxTable(elem, doc)
                if len(tbl.rows) < 2:
                    continue
                table_data = [_get_unique_row_cells(r) for r in tbl.rows]
                a, v = _parse_table_data(table_data)
                if a or v:
                    return "ok", a, v
                return "null", "", ""

        if not found_heading:
            # 未找到修订记录：字符量极少则视为仅封面页，否则为无修订记录
            if len(all_chars) < _COVER_PAGE_CHAR_THRESHOLD:
                return "cover_only", "", ""
            return "no_revision", "", ""

        return "null", "", ""

    except Exception as e:
        print(f"⚠️ docx 分析失败 [{os.path.basename(file_path)}]: {e}")
        return "null", "", ""


# ------------------------------------------------------------------
# 7b. PDF 综合分析（pdfplumber 文字层）
# ------------------------------------------------------------------
def _analyze_pdf(file_path: str) -> tuple[str, str, str]:
    """返回 (status, author, version)，含义同 _analyze_docx。"""
    try:
        import pdfplumber
    except ImportError:
        return "null", "", ""

    try:
        with pdfplumber.open(file_path) as pdf:
            if not pdf.pages:
                return "empty", "", ""

            # ── 步骤1：全文字量统计 ──────────────────────────────
            all_text = ''.join(page.extract_text() or '' for page in pdf.pages)
            all_chars = re.sub(r'\s', '', all_text)

            if len(all_chars) == 0:
                return "empty", "", ""

            # ── 步骤2：查找修订记录页（优先提取，成功则免判封面页）────
            # 注意：目录页也会含"修订记录"字样，需继续遍历找到真正的表格页
            found_revision_page = False
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if '修订记录' not in page_text:
                    continue
                found_revision_page = True
                for tbl in (page.extract_tables() or []):
                    if tbl and len(tbl) >= 2:
                        a, v = _parse_table_data(tbl)
                        if a or v:
                            return "ok", a, v
                # 当前页有"修订记录"但无有效表格，继续找下一页

            if not found_revision_page:
                # 未找到修订记录：页数极少且字少则视为仅封面页，否则为无修订记录
                if len(pdf.pages) <= 2 and len(all_chars) < _COVER_PAGE_CHAR_THRESHOLD:
                    return "cover_only", "", ""
                return "no_revision", "", ""

            return "null", "", ""

    except Exception as e:
        print(f"⚠️ PDF 分析失败 [{os.path.basename(file_path)}]: {e}")
        return "null", "", ""


# ------------------------------------------------------------------
# 7c. PDF OCR 分析（扫描版 PDF 兜底，使用 pymupdf 渲染，无需 poppler）
# ------------------------------------------------------------------
def _analyze_pdf_ocr(file_path: str) -> tuple[str, str, str]:
    """返回 (status, author, version)，含义同 _analyze_docx。"""
    try:
        import fitz          # pymupdf
        import pytesseract
        from PIL import Image
    except ImportError:
        return "null", "", ""

    try:
        doc = fitz.open(file_path)
        if len(doc) == 0:
            return "empty", "", ""

        # ── 步骤1：逐页渲染为图片并 OCR ─────────────────────────
        page_texts = []
        mat = fitz.Matrix(2, 2)   # 2× 缩放提高识别精度
        for page in doc:
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_texts.append(pytesseract.image_to_string(img, lang='chi_sim+eng'))

        all_chars  = re.sub(r'\s', '', ''.join(page_texts))

        if len(all_chars) == 0:
            return "empty", "", ""

        # ── 步骤2：查找修订记录页（优先提取，成功则免判封面页）────
        # 目录页也会含"修订记录"字样，需继续扫描找到真正的表格页
        found_revision_page = False
        for text in page_texts:
            if '修订记录' not in text:
                continue
            found_revision_page = True

            lines = [ln for ln in text.splitlines() if ln.strip()]
            header_idx = next(
                (i for i, ln in enumerate(lines)
                 if any(k in ln for k in _VERSION_KEYWORDS + _AUTHOR_KEYWORDS)),
                -1
            )
            if header_idx < 0:
                continue  # 未找到表头，继续下一页

            header     = re.split(r'\s{2,}|\t', lines[header_idx].strip())
            data_lines = [ln for ln in lines[header_idx + 1:] if ln.strip()]
            if not data_lines:
                continue  # 无数据行，继续下一页

            last_parts  = re.split(r'\s{2,}|\t', data_lines[-1].strip())
            version_col = author_col = -1
            for j, h in enumerate(header):
                if version_col < 0 and any(k in h for k in _VERSION_KEYWORDS):
                    version_col = j
                if author_col < 0 and any(k in h for k in _AUTHOR_KEYWORDS):
                    author_col = j

            version = last_parts[version_col] if 0 <= version_col < len(last_parts) else ""
            author  = last_parts[author_col]  if 0 <= author_col  < len(last_parts) else ""
            if author or version:
                return "ok", author, version
            # 找到表头和数据但列识别失败，继续下一页

        if not found_revision_page:
            # 未找到修订记录：页数极少且字少则视为仅封面页，否则为无修订记录
            if len(page_texts) <= 2 and len(all_chars) < _COVER_PAGE_CHAR_THRESHOLD:
                return "cover_only", "", ""
            return "no_revision", "", ""

        return "null", "", ""

    except Exception as e:
        print(f"⚠️ OCR 分析失败 [{os.path.basename(file_path)}]: {e}")
        return "null", "", ""


# ------------------------------------------------------------------
# 7d. 统一入口
# ------------------------------------------------------------------
def extract_author_version(file_path: str) -> str:
    """
    返回值规则：
      提取成功          → "张三  V2.0"
      文件为空          → "原文件内容为空"
      仅封面页          → "原文件只有封面页"
      有内容但无修订记录 → "原文件无修订记录"
      有修订记录但提取失败→ "原文件为空"
    """
    if not file_path or not os.path.exists(file_path):
        return "原文件为空"

    if os.path.getsize(file_path) == 0:
        return "原文件内容为空"

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        status, author, version = _analyze_docx(file_path)

    elif ext == ".pdf":
        status, author, version = _analyze_pdf(file_path)
        if status != "ok":
            # 扫描版 PDF：pdfplumber 无文字层会返回 "empty"；
            # 文字层不完整时可能返回 "no_revision" 或 "null"。
            # 以上三种情况都需要 OCR 补充确认。
            print(f"  → pdfplumber:{status}，尝试 OCR [{os.path.basename(file_path)}]")
            ocr_status, ocr_a, ocr_v = _analyze_pdf_ocr(file_path)
            # OCR 成功 → 直接采用
            if ocr_status == "ok":
                status, author, version = ocr_status, ocr_a, ocr_v
            # pdfplumber 认为是空文件或仅封面页（无文字层），但 OCR 读出了内容 → 以 OCR 为准
            elif status in ("empty", "cover_only") and ocr_status not in ("empty", "cover_only"):
                status, author, version = ocr_status, ocr_a, ocr_v

    else:
        print(f"⚠️ 不支持的格式 {ext}: {os.path.basename(file_path)}")
        return "原文件为空"

    if status == "empty":
        return "原文件内容为空"
    if status == "cover_only":
        return "原文件只有封面页"
    if status == "no_revision":
        return "原文件无修订记录"
    if status == "ok":
        return f"{author}  {version}" if (author and version) else (author or version or "原文件为空")
    return "原文件为空"


# =========================
# 8. 主循环：读取已下载文件，提取作者+版本
# =========================
results = []   # list of (title, owner, author_version)

for url, title_hex, owner in rows:
    title = decrypt_title(title_hex) if title_hex else ""
    title = normalize_plain_title(title)

    file_path = download_file(url, download_dir, title)
    author_version = extract_author_version(file_path) if file_path else "原文件为空"

    display_title = title if title else f"[解密失败] {title_hex[:24]}"
    results.append((display_title, owner, author_version))
    print(f"  [{display_title[:30]}]  作者/版本: {author_version}")

print(f"\n共处理 {len(results)} 条记录")

# =========================
# 9. 写入 Excel
# =========================
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "PR列表"

header_font = Font(bold=True, color="FFFFFF")
header_fill = PatternFill(fill_type="solid", fgColor="4472C4")
center      = Alignment(horizontal="center", vertical="center", wrap_text=True)

for col, header in enumerate(["标题", "负责人", "作者（当前版本）"], start=1):
    cell = ws.cell(row=1, column=col, value=header)
    cell.font      = header_font
    cell.fill      = header_fill
    cell.alignment = center

for row_idx, (title, owner, author_version) in enumerate(results, start=2):
    ws.cell(row=row_idx, column=1, value=title).alignment = Alignment(
        wrap_text=True, vertical="center"
    )
    ws.cell(row=row_idx, column=2, value=owner).alignment = Alignment(
        horizontal="center", vertical="center"
    )
    ws.cell(row=row_idx, column=3, value=author_version).alignment = Alignment(
        horizontal="center", vertical="center"
    )

ws.column_dimensions["A"].width = 60
ws.column_dimensions["B"].width = 15
ws.column_dimensions["C"].width = 25
ws.row_dimensions[1].height = 22

try:
    wb.save(output_path)
    print(f"已生成：{output_path}")
except PermissionError:
    print(f"\n❌ 保存失败：文件被占用，请先关闭 Excel 中已打开的 title_PR_output.xlsx，然后重新运行脚本。")
    print(f"   目标路径：{output_path}")
